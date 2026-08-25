"""DOCX rendering (task T098, research.md R12).

Rendered from the validated section model — never from HTML or a re-serialisation — so
what is exported is exactly what the RM approved. The exported content hash must match
the approval record, and it cannot if the export path re-derives content by another
route.

Every exported document carries three things the bank needs downstream and the RM
should not have to remember to add:

  * the approval record (who, when, and the exact content hash);
  * the Shariah review status, which is `PENDING_REVIEW` unless the Shariah function
    has acted — this system never clears it;
  * AI-assisted attribution (FR-036).
"""

from __future__ import annotations

import io

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from app.documents.models import ApprovalRecord, Document, DocumentVersion

# Warba green, used sparingly for headings.
_BRAND = RGBColor(0x00, 0x58, 0x4A)
_MUTED = RGBColor(0x5B, 0x66, 0x72)
_WARN = RGBColor(0x9A, 0x5B, 0x00)


def render_docx(
    document: Document,
    version: DocumentVersion,
    *,
    client_name: str,
    client_reference: str,
    approval: ApprovalRecord | None,
) -> bytes:
    """Render an approved document to DOCX bytes."""
    docx = DocxDocument()

    _apply_base_style(docx)
    _write_header(docx, document, client_name, client_reference)
    _write_sections(docx, version)
    _write_footer(docx, document, version, approval)

    buffer = io.BytesIO()
    docx.save(buffer)
    return buffer.getvalue()


def _apply_base_style(docx: DocxDocument) -> None:
    style = docx.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _write_header(
    docx: DocxDocument, document: Document, client_name: str, client_reference: str
) -> None:
    title = docx.add_heading(_display_name(document), level=0)
    for run in title.runs:
        run.font.color.rgb = _BRAND

    meta = docx.add_paragraph()
    meta.add_run(f"{client_name}\n").bold = True
    meta.add_run(f"Client reference: {client_reference}\n").font.color.rgb = _MUTED
    meta.add_run(f"Prepared: {document.created_at:%d %B %Y}").font.color.rgb = _MUTED

    # The attribution banner sits at the top, before the content, so a downstream
    # reader knows what they are holding before they read a word of it.
    banner = docx.add_paragraph()
    run = banner.add_run(
        "This document was drafted with AI assistance and reviewed and approved by the "
        "Relationship Manager named at the end of this document. Every factual statement "
        "is traceable to a source recorded in the system's audit trail."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = _MUTED

    docx.add_paragraph()


def _write_sections(docx: DocxDocument, version: DocumentVersion) -> None:
    for section in sorted(version.sections, key=lambda s: s.ordinal):
        heading = docx.add_heading(section.title, level=2)
        for run in heading.runs:
            run.font.color.rgb = _BRAND

        if section.content:
            docx.add_paragraph(section.content)
        else:
            placeholder = docx.add_paragraph()
            run = placeholder.add_run("No information was available for this section.")
            run.italic = True
            run.font.color.rgb = _MUTED

        # Gaps are exported, not silently dropped. A reader downstream must be able to
        # see what the RM could not source — that absence is part of the record, and
        # hiding it at export would undo the whole point of marking it.
        for gap in section.gaps or []:
            paragraph = docx.add_paragraph()
            paragraph.paragraph_format.left_indent = Pt(18)
            if gap.get("resolved"):
                run = paragraph.add_run(
                    f"{gap['label']} — resolved: {gap.get('resolution_note', '')}"
                )
                run.font.color.rgb = _MUTED
            else:
                run = paragraph.add_run(f"{gap['label']} — acknowledged as unavailable")
                run.font.color.rgb = _WARN
            run.font.size = Pt(9)
            run.italic = True

        if section.contains_external_data:
            note = docx.add_paragraph()
            run = note.add_run("Contains information from an external, unverified source.")
            run.font.size = Pt(8)
            run.italic = True
            run.font.color.rgb = _MUTED


def _write_footer(
    docx: DocxDocument,
    document: Document,
    version: DocumentVersion,
    approval: ApprovalRecord | None,
) -> None:
    docx.add_page_break()

    heading = docx.add_heading("Approval and Provenance", level=2)
    for run in heading.runs:
        run.font.color.rgb = _BRAND

    table = docx.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"

    rows: list[tuple[str, str]] = []

    if approval:
        rows += [
            ("Approved by", f"{approval.approver_name} ({approval.approver_role})"),
            ("Approved at", f"{approval.approved_at:%d %B %Y, %H:%M} UTC"),
            ("Content reference", approval.content_hash),
        ]
        if approval.gaps_acknowledged:
            rows.append(
                (
                    "Gaps acknowledged",
                    "; ".join(
                        f"{g['section_key']}.{g['field']}: {g['note']}"
                        for g in approval.gaps_acknowledged
                    ),
                )
            )
    else:  # pragma: no cover - export is gated on APPROVED
        rows.append(("Approval", "Not approved"))

    rows += [
        ("Shariah review status", document.shariah_status.value.replace("_", " ").title()),
        ("Document version", str(version.version_number)),
        ("Template version", version.template_version),
        ("Prompt version", version.prompt_version),
        ("Model", version.model_id or "n/a"),
    ]

    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        for run in cells[0].paragraphs[0].runs:
            run.bold = True

    # Stated explicitly, because "PENDING_REVIEW" alone could be misread as an
    # administrative formality rather than an outstanding requirement.
    if document.shariah_status.value == "PENDING_REVIEW":
        note = docx.add_paragraph()
        run = note.add_run(
            "Shariah review is outstanding. This document has not been cleared by the "
            "Shariah function."
        )
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = _WARN

    synthetic = docx.add_paragraph()
    synthetic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = synthetic.add_run("DEMONSTRATION DOCUMENT — all client data in this system is synthetic.")
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = _WARN


def _display_name(document: Document) -> str:
    return document.document_type.value.replace("_", " ").title()
